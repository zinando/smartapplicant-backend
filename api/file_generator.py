from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from io import BytesIO
from .ai import call_gemini
import json
import re
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT
import os
from auth_user.models import MatchedResumeData
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from .template_layouts import tp_layouts
from django.contrib.auth import get_user_model
from typing import Dict, List, Union



User = get_user_model()

def insert_paragraph_after(paragraph: Paragraph, text='', style=None):
    """
    Insert a new paragraph after the given paragraph.
    Formats skill categories as bold, followed by comma-separated skills.

    Args:
        paragraph (Paragraph): Reference paragraph
        text (str|dict): String or dict of {category: [skills]}
        style (str): Style to apply

    Returns:
        Paragraph: Newly inserted paragraph
    """
    # Get the XML element of the reference paragraph
    p_element = paragraph._p

    # Create new paragraph element
    new_p_element = OxmlElement('w:p')
    p_element.addnext(new_p_element)
    new_paragraph = Paragraph(new_p_element, paragraph._parent)

    # Handle dictionary of categorized skills
    if isinstance(text, dict):
        content = []
        for category, skills in text.items():
            # Add bold category name
            content.append(f"{category}: ")
            # Add comma-separated skills
            content.append(", ".join(skills))
        
        # Add all content to the paragraph
        run = new_paragraph.add_run("".join(content))
        
        # Make category names bold
        for category in text.keys():
            start = run.text.find(category)
            if start >= 0:
                end = start + len(category)
                run.bold = True  # Make the entire run bold
                break
        
        if style:
            new_paragraph.style = style
        return new_paragraph

    # Original string handling
    if text:
        new_paragraph.add_run(text)
    if style:
        new_paragraph.style = style

    return new_paragraph

def enforce_font(doc, font_name="Arial"):
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # Necessary for compatibility with Word

def add_contact_symbol(key, value):
    """Add contact symbol to the value based on the key."""
    symbols = {
        'email': '✉️',
        'phone': '📞',
        'git': '🐱',
        'linkedin': '🔗' if 'https' in value else '🔗 https://linkedin.com/in/',
        'portfolio': '🌐'
    }
    return f"{symbols.get(key, '')} {value}" if value else ''

class ResumeGenerator:
    template_path = "templates/general.docx"
    output_path = "templates/processed.docx"
    _document = None
    user_data = {
        "name": "John Doe",
        "professional_title": "Software Engineer",
        "city": "Ibadan",
        "state": "Oyo",
        "country": "Nigeria",
        "phone": "+2348023456789",
        "email": "jondoe@example.com",
        "git": "https://github.com/jondo",
        'linkedin': "https://linkedin.com/in/jon-do",
        "portfolio": "https://jondo.com",
        "objective": "To leverage my skills in software development to contribute to innovative projects.",
        "skills": ["Python", "JavaScript"],
        "certification": {"cert_name": "AWS Certified Solutions Architect", "cert_issuer": "Amazon", "issue_date": "2023-01", "cert_expiry": "2025-03 or Never"},
        "experience": {"company": "Tech Company", "position": "Software Engineer", "experience_duration": "2022-01 to 2023-01", "description": ["Developed web applications using Python and Django.", "Collaborated with cross-functional teams to define, design, and ship new features."], "achievements": ["Increased application performance by 20%.", "Led a team of 5 developers to deliver a project ahead of schedule."]}, 
        "project": {"title": "Automated Machine Monitoring System", "technologies": "python, pylogix, SQLite", "date": "2023", "description": ["Developed a real-time monitoring app to track PLC machine status and log downtimes.", "Improved maintenance response time by 40%."]},
        "education": {"degree": "BSc in Computer Science", "institution": "University of Ibadan", "location": "Oyo State Nigeria", "graduation_date": "2022-11", "description": "Graduated with First Class Honours"},
    }

    matching_user_data = {
        "name": "John Doe",
        "professional_title": "Software Engineer",
        "city": "Ibadan",
        "state": "Oyo",
        "country": "Nigeria",
        "git": "https://github.com/jondo",
        'linkedin': "https://linkedin.com/in/jon-do",
        "portfolio": "https://jondo.com",
        "phone": "+2348023456789",
        "email": "jondoe@example.com",
        "objective": "To leverage my skills in software development to contribute to innovative projects.",
        "skills": ["Python", "JavaScript"],
        "certifications": [
            {
                "cert_name": "AWS Certified Solutions Architect", 
                "cert_issuer": "Amazon", 
                "issue_date": "2023-01", 
                "cert_expiry": "2025-03 or Never",
                "description": "Demonstrated ability to design distributed systems on AWS"
            }, 
            {
                "cert_name": "Certified data analyst", 
                "cert_issuer": "IBM", 
                "issue_date": "2023-06", 
                "cert_expiry": "2025-03 or Never",
                "description": "Demonstrated ability to design distributed systems on AWS"
            }
        ],
        "experience": [
            {
                "company": "Tech Company", 
                "position": "Software Engineer", 
                "experience_duration": "2022-01 to 2023-01",
                "start_date": "2022-01",
                "end_date": "2023-01", 
                "description": ["Developed web applications using Python and Django.", "Collaborated with cross-functional teams to define, design, and ship new features."], 
                "achievements": ["Increased application performance by 20%.", "Led a team of 5 developers to deliver a project ahead of schedule."]
            }, 
            {...}
        ], 
        "education": [
            {
                "degree": "BSc", 
                "institution": "University of Ibadan",
                "location": "Oyo State Nigeria",
                "field_of_study": "Computer Science",
                "start_date": "2018-10",
                "end_date": "2022-11",
                "description": "Graduated with First Class Honours", 
                "graduation_date": "2022-11"
            }, 
            {...}
        ],
        "subjects": [
            {"title": "Educational Psychology & Classroom Management", "description": "A theory-and-practice course focusing on learner motivation, behavioral analysis, and inclusive education."},
            {"title": "Curriculum Design & Instructional Strategies", "description": "A course on designing effective curricula and employing diverse instructional strategies to enhance learning outcomes."},
            {"title": "Basic Technology", "description": "An introductory course on technology integration in education, covering digital tools and resources."},
            {...}
        ],
        "publications": [
            {"title": "Reforming Procurement Law in Nigeria: A Legislative Framework for Transparency", "source": "Published in the African Journal of Law & Public Policy, Vol. 12, Issue 3", "description": "Explores weaknesses in Nigeria’s Public Procurement Act and proposes amendments for better compliance and oversight.", "date": "2021"},
            {"title": "The Impact of E-Government on Public Service Delivery in Nigeria", "source": "Published in the Journal of Public Administration and Governance, Vol. 10, Issue 2", "description": "Analyzes the effectiveness of e-government initiatives in improving public service delivery in Nigeria.", "date": "2022"},
            {...}
        ],
        "teaching": [
            {"institution": "Federal University of Agriculture, Abeokuta", "role": "course instructor", "subject": "Molecular Genetics (BIO 304)", "description": "Taught 3rd-year undergraduates core concepts of molecular biology and genetic engineering", "date": "2023"},
            {"institution": "University of Ibadan", "role": "Teaching Assistant", "subject": "Introduction to Computer Science (CSC 101)", "description": "Assisted in teaching introductory computer science concepts to first-year students.", "date": "2022"},
            {...}
        ],
        "research": [
            {"title": "Investigating the Effects of Climate Change on Agricultural Productivity in Nigeria", "description": "A comprehensive study analyzing the impact of climate change on crop yields and food security in Nigeria.", "date": "2023"},
            {"title": "The Role of Renewable Energy in Sustainable Development", "description": "Explores how renewable energy sources can contribute to sustainable development goals in developing countries.", "date": "2022"},
            {...}
        ],
        "project": [
            {
                "title": "Automated Machine Monitoring System",
                "technologies": "python, pylogix, SQLite",
                "date": "2023",
                "description": ["Developed a real-time monitoring app to track PLC machine status and log downtimes.", "Improved maintenance response time by 40%."]
            },
            {...}
        ]
    }

    def __init__(self, resume_data=None, filename=None, matching=False, premium=True, user_id=0):
        self.filename = filename
        self.resume_data = resume_data
        self.user_object = User.objects.filter(id=user_id).first() if user_id != 0 else None
        if not matching:
            if premium:
                self.modify_matching_user_data()
                prompt = self.__generate_premium_prompt()
                self.matching_user_data = self.__promptGemini(prompt)
                self.deduct_resume_credit()
                self.log_matched_resume_data()
            else:
                prompt = self.__generate_prompt()
                self.user_data = self.__promptGemini(prompt)
        else:
            self.job_description = resume_data['job_description']
            prompt = self.__generate_matching_prompt()
            self.matching_user_data = self.__promptGemini(prompt)
            self.log_matched_resume_data()
            self.deduct_resume_credit()
    
    # method that subtracts resume credit from non-premium users after resume generation
    def deduct_resume_credit(self):
        """Deducts resume credit from non-premium users after resume generation."""
        if self.user_object and not self.user_object.is_premium:
            if self.user_object.resume_credits > 0:
                self.user_object.resume_credits -= 1
                self.user_object.save()
            else:
                raise ValueError("Insufficient resume credits.")
        
    def log_matched_resume_data(self):
        """Logs the matched resume data to the database."""
        if not self.matching_user_data:
            return
        
        # Check if user exists
        if not self.user_object:
            raise ValueError("User object is required to log matched resume data.")
        
        # Always save the matched resume data
        MatchedResumeData.objects.create(
            user=self.user_object,
            job_title=self.matching_user_data.get('professional_title', 'Unknown'),
            resume_data=self.matching_user_data
        )
    
    def modify_matching_user_data(self):
        """Modify the matching_user_data to ensure it meets the expected structure."""
        self.matching_user_data['education'][0]['graduation_date'] = '2023-08'
        # remove start_date and end_date from education
        self.matching_user_data['education'][0].pop('start_date', None)
        self.matching_user_data['education'][0].pop('end_date', None)
     
    def __generate_prompt(self) -> str:
        if self.resume_data is None:
            return "No resume data provided."
        return f"""
        Use the provided resume_data to generate ATS-friendly structured JSON.  
        Rules:  
        - Objective: craft a professional, keyword-rich statement aligned with the job title.  
        • If experienced → highlight expertise + value.  
        • If entry/transitioning → highlight eagerness to learn + contribute.  
        - Skills: must not be empty; only short items (≤35 chars). Include technical + industry soft skills. 4 skills max.  
        - Certifications: return only the most relevant one; if none, return null. If no expiry → "N/A".  
        - Experience: return only the most relevant job.  
        • Format: company, position, duration, description (list of strings), achievements (list of strings).  
        • At least 4 quantified achievements (use realistic values, e.g., 15%, 20%).  
        • If no end_date → "Present".  
        - Education: return only the most relevant education.  
        - All personal info (name, phone, email) must be clean + ATS-friendly.

        Resume_data:
        {self.resume_data}

        Expected JSON format:
        {self.user_data}           
        """
    def __generate_premium_prompt(self) -> str:
        if self.resume_data is None:
            return "No resume data provided."
        return f"""
        Use the provided resume_data to generate ATS-friendly structured JSON.  
        Rules:  
        - Objective: craft a professional, keyword-rich statement aligned with the job title.  
        • If experienced → highlight expertise + value.  
        • If entry/transitioning → highlight eagerness to learn + contribute.  
        - Skills: must not be empty; only short items (≤35 chars). Include technical + industry soft skills. 6 skills max.  
        - Certifications: return the most relevant ones; if none, return null. If no expiry → "N/A".  
        - Experience: return the most relevant jobs.  
        • Format: company, position, duration, description (list of strings), achievements (list of strings).  
        • At least 4 quantified achievements per job (use realistic values, e.g., 15%, 20%).  
        • If no end_date → "Present".  
        - Education: return the most relevant education entries.  
        - All personal info (name, phone, email) must be clean + ATS-friendly.

        Resume_data:
        {self.resume_data}

        Expected JSON format:
        {self.matching_user_data}           
        """
    
    def __generate_matching_prompt(self) -> str:
        # print(f"Resume text:\n{self.resume_data['resume_text']}")
        return f"""
        You are a professional resume editor and recruiter. 
        Your task is to re-write a candidate’s resume to make it closely match a specific job description (JD), optimizing it for maximum recruiter appeal and applicant tracking system (ATS) compatibility.

        Instructions:

        Analyze the job description and resume provided.

        Rewrite the resume to:

        Emphasize keywords and phrases found in the JD.

        Reposition and highlight relevant experiences, skills, and achievements.

        Remove irrelevant or unnecessary information that does not contribute to the JD.

        Maintain factual accuracy — do not fabricate sensitive information such as:

            Education

            Certifications

            Work experience

            Technical skills

            Projects

            You may infer and include missing but commonly expected content such as:

                Soft skills

                Summary/objective

                Role-aligned phrasing
        
        SPECIAL CONSIDERATIONS:

        
        * Ensure the information is presented in a clean, ATS-friendly format, especially the personal information like name, phone number and email.

        Work Experience Section:
        * if no descriptions are provided, generate a maximum of 3 based on the role
        * if no achievements is provided, generate possible imaginary ones based on the descriptions
        * return a maximum of 4 achievements items.
        * do not use place holders for quantification. use realistic imaginary quantification numbers like 15%, 20%, 30%, etc.

        Education Section:
        * if no descriptions are provided, return imaginary possible one
        * graduation date is the same as end_date
        * the degree value should be the degree abreviation: B.Sc, M.Sc, OND, HND etc

        Certification Section:
        * if no expiry date is provide, return N/A
        * if no issue date is provided, return in-progress
        * if no description is provided, return imaginary one

        Subjects Section:
        * if no description is provided, return imaginary possible one
        * this section is for teachers and educators only, so it should contain subjects they have taught or are qualified to teach.

        Publications Section:
        * if no description is provided, return imaginary possible one

        Teaching Section:
        * if no description is provided, return imaginary possible one
        * this section is for science and research lecturers, so it should contain subjects they have taught or are qualified to teach.
        * it should be different from the subjects section, which is for teachers and educators, and work experience section, which is for general professionals in other fields.
        
        Research Section:
        * if no description is provided, return imaginary possible one
        * this section is for science and research lecturers, so it should contain research they have interest in.  


        Skills Section:
        * each item should be a single word or a short phrase not more than 35 characters long.
        For a ui/ux designer, return skills as a list of dictionaries with the following keys: Design, Prototyping, Research, Project Tools, Soft Skills, Branding, Code. Their respective values should be a list of skills relevant to the key.

        We will use your data to generate a clean, modern, recruiter-friendly resume with clearly labeled sections (e.g., Summary, Skills, Experience, Education).

        If any of the following are missing — job description, resume, or structural format — return an empty string.

        Job Description (JD):
        {self.job_description}

        Resume_data:
        {self.resume_data['resume_text']}

        Expected JSON format:
        {self.matching_user_data}
        """
    
    def __promptGemini(self, prompt: str) -> dict:
        # Placeholder for the actual Gemini API call
        # print("Prompting Gemini with:", prompt)
        try:
            text = call_gemini(prompt)
            # print(f'Raw Gemini response: {text}')
            text = text.replace('```json', '')
            text = text.replace('```', '')
            # print(f'Gemini response: {text}')
            response = json.loads(text)
        except Exception as e:
            print(f"Error calling Gemini: {e}")
            response = {}
        
        return response

    def _insert_bullets(self, doc, anchor_text, items, bullet=True):
        for i, para in enumerate(doc.paragraphs):
            if anchor_text in para.text:
                para.text = para.text.replace(anchor_text, "")
                if isinstance(items, dict):
                    for key, value in items.items():
                        mr = {}
                        mr[key] = value
                        # Insert bullet paragraph after the anchor
                        insert_paragraph_after(para, mr, style='List Bullet')
                elif isinstance(items, list):
                    for item in items:
                        # Insert bullet paragraph starting from the anchor
                        insert_paragraph_after(para, item, style='List Bullet' if bullet else None)
                        
                break
    
    def __select_template(self, template_id:str=''):
        """Select the appropriate template based on user data."""

        if template_id:
            self.template_path = f"templates/{template_id}.docx"        
        else:
            self.template_path = "templates/general_certification.docx"
        
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"{self.template_path} not found")
        
    def populate_template(self):
        """Populate a Word template with user data and save the result."""
        # Early exit checks
        if not self.user_data:
            print("No user data available.")
            return ''
        if not self.filename:
            print("No filename available.")
            return ''
        
        # Get template ID
        template_id = self.resume_data.get('template_id', '')

        if template_id == 'modern':
            try:
                self.__select_template(template_id)
                template_layout = tp_layouts.get(template_id)
                user_data = self.user_data.copy()

                # Use ColumnAwareTemplatePopulator for modern templates
                populator = ColumnAwareTemplatePopulator(self.template_path, template_layout)
                this_doc = populator.populate_template(user_data)
                if not this_doc:
                    raise Exception("Failed to populate modern template.")
                
                # Clean unused placeholders
                self._clean_unused_placeholders(this_doc)

                # Apply font style to the entire document
                enforce_font(this_doc, font_name="Arial")

                # save document to file
                self.__save_document(this_doc)
                return self.filename
            except Exception as e:
                print({str(e)})
                return ''

            
        # Select the template based on the provided template ID
        self.__select_template(template_id)

        try:
            doc = Document(self.template_path)
            placeholder_cache = {}  # Cache for parsed placeholders

            # Process paragraphs
            for para in doc.paragraphs:
                if not para.text:  # Skip empty paragraphs
                    continue
                    
                para_text = para.text
                modified = False
                
                # Check for simple placeholders first
                for key, value in self.user_data.items():
                    if isinstance(value, str):
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in para_text:
                            para_text = para_text.replace(placeholder, value)
                            # self._apply_style_to_run(para, key, value)
                            modified = True
                    elif isinstance(value, dict):
                        for sub_key, sub_value in value.items():
                            if isinstance(sub_value, str):
                                placeholder = f"{{{{{sub_key}}}}}"
                                if placeholder in para_text:
                                    para_text = para_text.replace(placeholder, sub_value)
                                    modified = True
                    elif isinstance(value, list):
                        # 1. Determine the list length
                        # 2. Get one of the list items
                        # 3. Make a list of placeholders based on the keys of the item
                        # 4. Check and isolate all the paragraphs that contain any of the placeholders in order
                        # 5. For each item in the list, duplicate the isolated paragraphs and replace the placeholders with the item values
                        # 6. Insert the duplicated paragraphs after the last isolated paragraph
                        
                        # Handle list of dicts (e.g., certifications, experience)
                        if all(isinstance(item, dict) for item in value):
                            # Get placeholders from the first item
                            first_item = value[0]
                            placeholders = [f"{{{{{k}}}}}" for k in first_item.keys()]
                            
                            # Find paragraphs containing any of the placeholders
                            target_paragraphs = []
                            for p in doc.paragraphs:
                                if any(ph in p.text for ph in placeholders):
                                    target_paragraphs.append(p)
                            
                            if target_paragraphs:
                                last_para = target_paragraphs[-1]
                                parent = last_para._element.getparent()
                                index = parent.index(last_para._element)
                                
                                # Remove original paragraphs
                                for p in target_paragraphs:
                                    parent.remove(p._element)
                                
                                # Insert new paragraphs for each item
                                for item in value:
                                    for template_para in target_paragraphs:
                                        new_para = insert_paragraph_after(last_para, style=template_para.style.name)
                                        new_para.text = template_para.text
                                        for k, v in item.items():
                                            if isinstance(v, str):
                                                ph = f"{{{{{k}}}}}"
                                                new_para.text = new_para.text.replace(ph, v)
                                        last_para = new_para
                                modified = True
                        # if all(isinstance(item, dict) for item in value):
                        #     for item in value:
                        #         for sub_key, sub_value in item.items():
                        #             if isinstance(sub_value, str):
                        #                 placeholder = f"{{{{{sub_key}}}}}"
                        #                 if placeholder in para_text:
                        #                     para_text = para_text.replace(placeholder, sub_value)
                        #                     modified = True
                
                # Update paragraph if modified
                if modified:
                    # Clear existing runs more efficiently
                    para.clear()
                    para.add_run(para_text)
                    if self.user_data.get('name') in para_text:
                        # Apply specific style for name (make the font size 28)
                        self._apply_style_to_run(para, 'name', self.user_data.get('name', ''))
                    

            # Process lists
            list_fields = {
                "{{skills_list}}": self.user_data.get("skills", []),
                "{{certifications_list}}": self.user_data.get("certifications", []),
                "{{description}}": self.user_data.get("experience", {}).get("description", []),
                "{{achievements}}": self.user_data.get("experience", {}).get("achievements", []),
            }

            for anchor, items in list_fields.items():
                print(f"Processing list for anchor {anchor} with items: {items}")
                if items:  # Only process if there are items
                    if anchor == "{{description}}":
                        self._insert_bullets(doc, anchor, items, bullet=False)
                    else:
                        self._insert_bullets(doc, anchor, items)
            
            # Remove unused placeholders
            self._clean_unused_placeholders(doc)

            self.__save_document(doc)
            
            return self.filename
            
        except Exception as e:
            print(f"Error populating template: {str(e)}")
            return ''
        
    def populate_premium_template(self):
        """Uses the ATS Bold Classic template style for premium users."""
        result = ''
        if self.resume_data.get('template_id', '') == 'professional':
            # change template id to premium template
            self.resume_data['template_id'] = 'professional_premium'

            # select premium template
            self.__select_template('professional_premium')
            result = self.populate_ats_bold_classic_resume()
        elif self.resume_data.get('template_id', '') == 'modern':
            # change template id to premium template
            self.resume_data['template_id'] = 'modern_premium'
            try:
                self.__select_template(self.resume_data.get('template_id', ''))
                template_layout = tp_layouts.get(self.resume_data.get('template_id', ''))
                user_data = self.matching_user_data.copy()

                # if experience has more than one item, create a new key 'experience_continuation' with the rest of the items
                if 'experience' in user_data and isinstance(user_data['experience'], list) and len(user_data['experience']) > 1:
                    user_data['experience_continuation'] = user_data['experience'][1:]
                    user_data['experience'] = user_data['experience'][0]

                # Use ColumnAwareTemplatePopulator for modern templates
                populator = ColumnAwareTemplatePopulator(self.template_path, template_layout)
                this_doc = populator.populate_template(user_data)
                if not this_doc:
                    raise Exception("Failed to populate modern template.")
                
                # Clean unused placeholders
                self._clean_unused_placeholders(this_doc)

                # Apply font style to the entire document
                enforce_font(this_doc, font_name="Arial")

                # save document to file
                self.__save_document(this_doc)
                result = self.filename
            except Exception as e:
                print({str(e)})
                result = ''


        if result == '':
            raise Exception('Failed to populate premium resume template')
        return result
        
    def populate_matching_template(self, template_id:str):
        """Populate a Word template with user data and save the result."""
        result = ''
        # Select template
        self.__select_template(template_id)

        if template_id == 'ats_bold_classic_resume' or 'ats_bold_classic_resume' in template_id:
            result = self.populate_ats_bold_classic_resume()
        
        if result == '':
            raise Exception('Failed to populate resume template')
        
        return result

    def populate_ats_bold_classic_resume(self):
        """ Processes the ATS Bold Classic template (v2.1)
        
            Requirements:
            - Requires experience in reverse-chronological order
            - Skills must be < 25 characters each
        """
        # helper functions
        def process_simple_fields(doc):
            """Process simple key-value pairs in the document."""
            for para in doc.paragraphs:
                if not para.text:
                    continue
                    
                para_text = para.text
                modified = False
                
                for key, value in self.matching_user_data.items():
                    if isinstance(value, str):
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in para_text:
                            para_text = para_text.replace(placeholder, value)
                            modified = True
                    elif isinstance(value, dict) and not any(k in key.lower() for k in ['education', 'skills', 'experience', 'certification', 'projects']):
                        for sub_key, sub_value in value.items():
                            if isinstance(sub_value, str):
                                placeholder = f"{{{{{sub_key}}}}}"
                                if placeholder in para_text:
                                    para_text = para_text.replace(placeholder, sub_value)
                                    modified = True
                
                if modified:
                    para.clear()
                    para.add_run(para_text)
                    if self.matching_user_data.get('name') in para_text:
                        # Apply specific style for name (make the font size 28)
                        self._apply_style_to_run(para, 'name', self.user_data.get('name', ''))
        
        def process_multi_item_sections(doc):
            """Process sections with perfectly aligned dates at line ends."""
            
            template_id = self.resume_data.get('template_id')
            multi_item_sections = tp_layouts.get(template_id)

            if not multi_item_sections:
                multi_item_sections = tp_layouts['ats_bold_classic_resume']
                
            for section, config in multi_item_sections.items():
                items = self.matching_user_data.get(section, [])
                if not items:
                    continue
                    
                section_paragraphs = []
                
                for item_idx, item in enumerate(items):
                    item_paragraphs = []
                    
                    for line_config in config['item_template']:
                        # Create new paragraph if needed
                        if line_config.get('new_paragraph', True):
                            p = doc.add_paragraph()
                            item_paragraphs.append(p)
                            
                            # Configure tab stops for dated lines
                            if line_config.get('type') == 'dated_line':
                                tab_stops = p.paragraph_format.tab_stops
                                tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
                            
                            if line_config.get('bullet', False):
                                p.style = 'List Bullet'
                                if line_config.get('indent', 0) > 0:
                                    p.paragraph_format.left_indent = Inches(line_config['indent'])
                                    p.paragraph_format.first_line_indent = Inches(-0.25)  # Proper bullet spacing
                        
                        # Handle dated lines with perfect date alignment
                        if line_config.get('type') == 'dated_line':
                            content = line_config['content']
                            dates = line_config['dates']
                            
                            # Replace placeholders in content
                            for key, value in item.items():
                                if isinstance(value, str):
                                    placeholder = f"{{{{{key}}}}}"
                                    content = content.replace(placeholder, value)
                                    dates = dates.replace(placeholder, value)
                            
                            # Add content and dates with tab separation
                            run = p.add_run(content)
                            if line_config.get('content_bold', False):
                                run.bold = True
                            p.add_run('\t')  # Tab character for alignment
                            p.add_run(dates)
                            continue
                        
                        # Handle list-type descriptions/achievements
                        if line_config.get('is_list', False) and line_config.get('content_key', '') in item and isinstance(item[line_config.get('content_key', '')], list):
                            for i, bullet_point in enumerate(item[line_config.get('content_key', '')]):
                                if i > 0 and line_config.get('bullet', True):
                                    p = doc.add_paragraph(style='List Bullet')
                                    p.paragraph_format.left_indent = Inches(line_config.get('indent', 0.5))
                                    p.paragraph_format.first_line_indent = Inches(-0.25)
                                    item_paragraphs.append(p)
                                
                                p.add_run(bullet_point)
                            continue
                        
                        # Handle regular text
                        if 'text' in line_config:
                            line_text = line_config['text']
                            for key, value in item.items():
                                if isinstance(value, str):
                                    placeholder = f"{{{{{key}}}}}"
                                    line_text = line_text.replace(placeholder, value)
                                elif isinstance(value, list):
                                    value = "\n".join(value)
                                    placeholder = f"{{{{{key}}}}}"
                                    line_text = line_text.replace(placeholder, value)
                            
                            run = p.add_run(line_text)
                            if line_config.get('bold', False) or line_config.get('content_bold', False):
                                run.bold = True
                            if line_config.get('italic', False):
                                run.italic = True
                    
                    # Add separator between items (except after last item)
                    if item_idx < len(items) - 1:
                        separator_para = doc.add_paragraph()
                        separator_para.add_run(config['separator'])
                        item_paragraphs.append(separator_para)
                    
                    section_paragraphs.extend(item_paragraphs)
                
                # Replace placeholder with styled content
                replace_placeholder_with_styled_content(
                    doc,
                    config.get('template_placeholder') or config.get('template_anchor'),
                    section_paragraphs
                )

        def process_bullet_lists(doc):
            """Process simple bullet list sections."""
            list_fields = {
                "{{skills_list}}": self.matching_user_data.get("skills", []),
            }

            for anchor, items in list_fields.items():
                if items:
                    self._insert_bullets(doc, anchor, items)

        def replace_placeholder_with_styled_content(doc, placeholder, new_paragraphs):
            """Replace a placeholder with styled paragraphs."""
            for para in doc.paragraphs:
                if placeholder in para.text:
                    # Get the parent element of the paragraph
                    parent = para._element.getparent()
                    # Get the index of the paragraph in its parent
                    index = parent.index(para._element)
                    
                    # Remove the placeholder paragraph
                    parent.remove(para._element)
                    
                    # Insert the new paragraphs at the same position
                    for new_para in reversed(new_paragraphs):
                        parent.insert(index, new_para._element)
                    
                    break

        # Early exit checks
        if not self.matching_user_data:
            # print("No user data available.")
            return ''
        if not self.filename:
            # print("No filename available.")
            return ''

        try:
            doc = Document(self.template_path)
            self._document = doc

            # Process simple fields (name, contact info, etc.)
            process_simple_fields(doc)
            # print("Simple fields processed successfully.")
            
            # Process multi-item sections
            process_multi_item_sections(doc)
            # print("Multi-item sections processed successfully.")
            
            # Process bullet lists
            process_bullet_lists(doc)
            # print("Bullet lists processed successfully.")

            # Remove unused placeholders
            self._clean_unused_placeholders(doc)
            # print("Unused placeholders cleaned successfully.")

            # Save the document
            os.makedirs('generated_docs', exist_ok=True)
            save_path = os.path.join('generated_docs', self.filename)
            doc.save(save_path)
            
            return self.filename
            
        except Exception as e:
            print(f"Error populating template: {str(e)}")
            return ''
    
    def _clean_unused_placeholders(self, doc: Document) -> None:
        """
        Remove all unused placeholders and their associated formatting elements.
        Handles three cases:
        1. Simple inline placeholders (remove placeholder only)
        2. Labeled placeholders (remove label + placeholder)
        3. Sectional placeholders (remove section header + divider + placeholder)
        """
        # Define placeholder patterns for each case
        placeholder_groups = {
            'simple': [
                '{{name}}', '{{email}}', '{{phone}}',
                '{{city}}', '{{state}}', '{{country}}', '{{experience_section_continuation}}',
                '{{git}}', '{{portfolio}}', '{{linkedin}}', '{{description}}', '{{achievements}}',
            ],
            'labeled': [
                ('Git:', '{{git}}'),
                ('Portfolio:', '{{portfolio}}'), 
                ('LinkedIn:', '{{linkedin}}')
            ],
            'sectional': [
                '{{certification_section}}', '{{experience_section}}', '{{education_section}}',
                '{{project_section}}', '{{skils_list}}', '{{skills_section}}', '{{cert_name}}', '{{cert_issuer}}', '{{cert_issue_date}}', '{{cert_expiry}}',
                '{{publication_section}}', '{{objective}}', '{{subject_section}}', '{{teaching_section}}', '{{research_interest_section}}'
            ]
        }

        # Process sectional placeholders
        for placeholder in placeholder_groups['sectional']:
            self._remove_sectional_placeholder(doc, placeholder)

         # Process labeled placeholders
        for label, placeholder in placeholder_groups['labeled']:
            self._remove_labeled_placeholder(doc, label, placeholder)

        # Process simple placeholders
        for placeholder in placeholder_groups['simple']:
            self._remove_simple_placeholder(doc, placeholder)

    def _remove_simple_placeholder(self, doc: Document, placeholder: str) -> None:
        """Remove simple inline placeholders that weren't replaced."""
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, "").strip()
                # Remove paragraph if empty
                if not paragraph.text:
                    self._remove_paragraph(paragraph)

    def _remove_labeled_placeholder(self, doc: Document, label: str, placeholder: str) -> None:
        """
        Remove placeholders with preceding labels.
        Example: "Git: {{git}}" -> removes entire line when placeholder exists
        """
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text and label in paragraph.text:
                # Check if this is the only content in the paragraph
                full_pattern = f"{label} {placeholder}"
                if paragraph.text.strip() == full_pattern:
                    self._remove_paragraph(paragraph)
                else:
                    # Handle cases where label+placeholder are part of larger text
                    paragraph.text = paragraph.text.replace(full_pattern, "").strip()

    def _remove_sectional_placeholder(self, doc: Document, placeholder: str) -> None:
        """
        Remove sectional placeholders and their associated header structure.
        Removes:
        - The paragraph containing the placeholder
        - The divider paragraph immediately above it (index-1)
        - The section header paragraph above that (index-2)
        """
        paragraphs = doc.paragraphs
        to_remove = set()

        for i, paragraph in enumerate(paragraphs):
            if placeholder in paragraph.text:
                # print('sectionl placeholder found:', placeholder)
                # Always remove the placeholder paragraph
                to_remove.add(i)
                
                # Remove divider paragraph (immediately above)
                if i > 0:
                    to_remove.add(i-1)
                    
                    # Remove section header (two above)
                    if i > 1:
                        to_remove.add(i-2)

        # Remove marked paragraphs in reverse order to maintain indices
        for i in sorted(to_remove, reverse=True):
            if i < len(paragraphs):
                self._remove_paragraph(paragraphs[i])

    def _is_horizontal_rule(self, paragraph: Paragraph) -> bool:
        """More reliable horizontal rule detection."""
        text = paragraph.text.strip()
        
        # Check for line characters
        line_chars = {'─', '━', '═', '—', '―', '-', '_', '='}
        if text and all(c in line_chars for c in text):
            return True
            
        # Check for empty paragraph with special formatting
        if not text and paragraph.paragraph_format.space_after == 0:
            return True
            
        return False

    def _remove_paragraph(self, paragraph: Paragraph) -> None:
        """Safely remove a paragraph from the document."""
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    
    def __save_document(self, doc: Document) -> None:
        """Save the populated document to the specified filename."""
        os.makedirs('generated_docs', exist_ok=True)
        save_path = os.path.join('generated_docs', self.filename)
        doc.save(save_path)
    
    def _apply_style_to_run(self, para, key, value):
        """Optional styling based on field type."""
        for run in para.runs:
            if key == 'name':
                run.bold = True
                run.font.size = Pt(28)
            elif key == 'professional_title':
                run.italic = True
                run.font.size = Pt(11)
            elif key in ['email', 'phone']:
                run.font.size = Pt(10)

class ColumnAwareTemplatePopulator:
    def __init__(self, template_path: str, layout_config: Dict):
        self.doc = Document(template_path)
        self.layout_config = layout_config

    def populate_template(self, user_data: Dict) -> Document:
        self.user_data = user_data
        self._process_inline_fields()
        self._process_multi_item_sections()
        return self.doc

    def _process_inline_fields(self):
        for para in self.doc.paragraphs:
            if not para.text:
                continue
            text = para.text
            modified = False

            if isinstance(self.user_data, dict):
                for key, value in self.user_data.items():
                    if isinstance(value, str):
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in text:
                            # add contact symbol
                            value = add_contact_symbol(key, value)
                            text = text.replace(placeholder, value)
                            modified = True
                    elif isinstance(value, dict):
                        for sub_key, sub_val in value.items():
                            if isinstance(sub_val, str):
                                placeholder = f"{{{{{sub_key}}}}}"
                                if placeholder in text:
                                    text = text.replace(placeholder, sub_val)
                                    modified = True

                if modified:
                    para.clear()
                    para.add_run(text)
                    if self.user_data.get('name') in text:
                        # Apply specific style for name (make the font size 28)
                        self._apply_style_to_run(para, 'name', self.user_data.get('name', ''))

    def _process_multi_item_sections(self):
        if isinstance(self.layout_config, dict):
            for section, config in self.layout_config.items():
                placeholder = config.get("template_anchor")
                if not placeholder:
                    continue

                items = self.user_data.get(section, [])
                if not items:
                    continue

                for para in self.doc.paragraphs:
                    if placeholder in para.text:
                        # print(f"Processing section: {section} with placeholder: {placeholder}")
                        new_paragraphs = []
                        for item_index, item in enumerate(items if isinstance(items, list) else [items]):
                            item_paragraphs = []

                            for line_config in config["item_template"]:
                                new_para = insert_paragraph_after(para, style="List Bullet" if line_config.get("bullet") else None)

                                # Handle list of strings (is_list)
                                if line_config.get("is_list"):
                                    key = self._extract_placeholder_key(line_config["text"])
                                    list_items = item.get(key, []) if isinstance(item, dict) else item
                                    if isinstance(list_items, list):
                                        for li in list_items:
                                            para_style = "List Bullet" if line_config.get("bullet", False) else None
                                            new_line = insert_paragraph_after(new_para, li, style=para_style)

                                            if line_config.get("bullet"):
                                                indent_val = line_config.get("indent", 0)
                                                if indent_val > 0:
                                                    new_line.paragraph_format.left_indent = Inches(indent_val)
                                                    new_line.paragraph_format.first_line_indent = Inches(-0.25)

                                            item_paragraphs.append(new_line)
                                        continue


                                # Handle dated line
                                if line_config.get("type") == "dated_line":
                                    content = self._replace_placeholders(line_config["content"], item)
                                    dates = self._replace_placeholders(line_config["dates"], item)

                                    new_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
                                    run = new_para.add_run(content)
                                    if line_config.get("content_bold"):
                                        run.bold = True
                                    new_para.add_run('\t')
                                    new_para.add_run(dates)
                                    item_paragraphs.append(new_para)
                                    continue

                                # Handle string or templated line
                                if "text" in line_config:
                                    line = self._replace_placeholders(line_config["text"], item)
                                    run = new_para.add_run(line)
                                    if line_config.get("content_bold"):
                                        run.bold = True
                                    if line_config.get("italic"):
                                        run.italic = True
                                    item_paragraphs.append(new_para)

                            # Optional separator
                            if item_index < len(items) - 1 and config.get("separator"):
                                sep_para = insert_paragraph_after(item_paragraphs[-1], config["separator"])
                                item_paragraphs.append(sep_para)

                            new_paragraphs.extend(item_paragraphs)

                        # Replace original placeholder paragraph
                        self._replace_placeholder_with_paragraphs(para, new_paragraphs)
                        break

    def _process_list_only_sections(self):
        if isinstance(self.layout_config, dict):
            for section, config in self.layout_config.items():
                if not config.get("item_template"):
                    continue

                # Only apply this logic if the section is flat (list of strings)
                items = self.user_data.get(section)
                if not isinstance(items, list) or not items:
                    continue

                placeholder = config.get("template_anchor")
                if not placeholder:
                    continue

                # We assume only one line_config for flat list sections
                line_config = config["item_template"][0]

                for para in self.doc.paragraphs:
                    if placeholder in para.text:
                        new_paragraphs = []

                        for item in items:
                            para_style = "List Bullet" if line_config.get("bullet", False) else None
                            bullet_para = insert_paragraph_after(para, item, style=para_style)

                            if line_config.get("bullet"):
                                indent_val = line_config.get("indent", 0)
                                if indent_val > 0:
                                    bullet_para.paragraph_format.left_indent = Inches(indent_val)
                                    bullet_para.paragraph_format.first_line_indent = Inches(-0.25)

                            new_paragraphs.append(bullet_para)

                        self._replace_placeholder_with_paragraphs(para, new_paragraphs)
                        break

    def _replace_placeholder_with_paragraphs(self, placeholder_para: Paragraph, new_paragraphs: List[Paragraph]):
        parent = placeholder_para._element.getparent()
        index = parent.index(placeholder_para._element)
        parent.remove(placeholder_para._element)
        for p in reversed(new_paragraphs):
            parent.insert(index, p._element)

    def _remove_unused_placeholders(self):
        for para in self.doc.paragraphs:
            if "{{" in para.text and "}}" in para.text:
                para.text = ""

    def _extract_placeholder_key(self, text: str) -> str:
        match = re.search(r"\{\{(\w+)\}\}", text)
        return match.group(1) if match else ""

    def _replace_placeholders(self, text: str, data: Union[Dict, str]) -> str:
        if isinstance(data, dict):
            # Find all placeholders in the text
            matches = re.findall(r"{{(.*?)}}", text)
            for key in matches:
                value = data.get(key, "")
                if isinstance(value, (str, int, float)):
                    text = text.replace(f"{{{{{key}}}}}", str(value))
                else:
                    text = text.replace(f"{{{{{key}}}}}", "")
        elif isinstance(data, str):
            # Simple case like skill strings
            text = text.replace("{{skill}}", data)
        return text
    
    def _apply_style_to_run(self, para, key, value):
        """Optional styling based on field type."""
        for run in para.runs:
            if key == 'name':
                run.bold = True
                run.font.size = Pt(28)
            elif key == 'professional_title':
                run.italic = True
                run.font.size = Pt(11)
            elif key in ['email', 'phone']:
                run.font.size = Pt(10)

